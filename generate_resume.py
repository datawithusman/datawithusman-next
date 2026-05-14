from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Page Setup - tight margins
for section in doc.sections:
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = Pt(12)

DARK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x1B, 0x3A, 0x5C)
GREY = RGBColor(0x55, 0x55, 0x55)

def section_head(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    # bottom border
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="2" w:space="1" w:color="1B3A5C"/></w:pBdr>')
    pPr.append(pBdr)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(9)
        rb.font.name = 'Calibri'
        rb.font.color.rgb = DARK
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

def job_line(company, location, title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(company)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    r2 = p.add_run(f" | {location} | {title} | {dates}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY
    r2.font.name = 'Calibri'

# ══════════════════════════════════════════
# HEADER - Single compact block
# ══════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
r = p.add_run("MUHAMMAD USMAN")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(1)
p2.paragraph_format.line_spacing = Pt(11)
r = p2.add_run("Backend Developer | Python Instructor | Data Analyst")
r.font.size = Pt(9.5)
r.font.color.rgb = GREY
r.font.name = 'Calibri'
r.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(2)
p3.paragraph_format.line_spacing = Pt(11)
r = p3.add_run("+966-56-027-6175 | datawithusman@gmail.com | linkedin.com/in/datawithusman | github.com/datawithusman | datawithusman.com")
r.font.size = Pt(8.5)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'

# ══════════════════════════════════════════
# SUMMARY - 2 lines
# ══════════════════════════════════════════
section_head("Summary")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(1)
p.paragraph_format.line_spacing = Pt(11)
r = p.add_run(
    "Results-driven Backend Developer and Python Instructor with 2+ years of experience in software development, "
    "data analytics, and operations. Currently building AI-powered backend systems at Nobel AI (Node.js) and teaching "
    "Python worldwide as a Stanford CIP Section Leader. Harvard CS50x certified with a 3.85 CGPA in BBIT."
)
r.font.size = Pt(9)
r.font.name = 'Calibri'

# ══════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════
section_head("Professional Experience")

job_line("Nobel AI", "Remote", "Backend Developer", "Mar 2026 - Present")
bullet("Developing AI-powered backend systems using Node.js for real-time soft skills analysis in virtual meetings; building scalable RESTful APIs and microservices.", "")

job_line("New Packaging Delivery Services", "Bahrain", "Operations Manager", "May 2023 - Apr 2025")
bullet("Managed KPIs for 200+ rider fleet for Talabat; built data-driven tracking system in Excel & Power BI optimizing route efficiency and supplier ratings.", "")

# ══════════════════════════════════════════
# TEACHING EXPERIENCE
# ══════════════════════════════════════════
section_head("Teaching Experience")

job_line("Stanford University (Code in Place)", "Remote", "Section Leader - Python", "Apr 2026 - Present")
bullet("Teaching Python fundamentals to students worldwide in an intensive 6-week program; leading live coding sessions and mentoring through debugging and problem-solving.", "")

job_line("Nobel AI", "Remote", "Python Bootcamp Instructor", "May - Jun 2026")
bullet("Designed and delivered Python bootcamp for worldwide beginners covering fundamentals through project-based learning over 6 weeks.", "")

job_line("iCodeGuru", "Remote", "Python Instructor", "2025 - Present")
bullet("Teach Python fundamentals (OOP, loops, functions) to beginners via live coding; mentor students on debugging and real-world projects.", "")

# ══════════════════════════════════════════
# KEY PROJECTS
# ══════════════════════════════════════════
section_head("Key Projects")

projects = [
    ("Nobel AI Core", "Node.js, AI/ML", "Real-time soft skills analysis engine | github.com/datawithusman/Nobel-AI-Core"),
    ("AI Trading Agent", "Python, ML", "lablab.ai hackathon - automated trading agent | github.com/datawithusman"),
    ("Telecom Churn Prediction", "Scikit-learn, Pandas", "ML churn prediction with business insights | github.com/datawithusman/telecom-churn-prediction"),
    ("Model Blind Spots Dataset", "AI Safety, NLP", "Published evaluation dataset on Hugging Face | huggingface.co/datasets/Datawithusman"),
    ("Easy Life Clinic System", "Full Stack", "Booking & management system for aesthetic clinic | github.com/datawithusman/easy-life-clinic"),
]

for title, tech, desc in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"{title}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    r2 = p.add_run(f" ({tech}) - ")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = GREY
    r2.font.name = 'Calibri'
    r2.italic = True
    r3 = p.add_run(desc)
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    r3.font.name = 'Calibri'

# ══════════════════════════════════════════
# TECHNICAL SKILLS - inline format
# ══════════════════════════════════════════
section_head("Technical Skills")

skill_lines = [
    ("Languages: ", "Python (Core, OOP, Data Structures), JavaScript, Node.js, SQL"),
    ("Frameworks & Tools: ", "Next.js, React, REST APIs, Git/GitHub, Power BI, Tableau, Excel (Power Query)"),
    ("Data & AI: ", "Scikit-learn, Pandas, NumPy, Machine Learning, KNIME, Data Visualization"),
    ("Operations: ", "Fleet Management, KPI Tracking, Workflow Optimization, Vendor Coordination"),
]

for bold_text, normal_text in skill_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    r = p.add_run(bold_text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    r.font.name = 'Calibri'
    r2 = p.add_run(normal_text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    r2.font.name = 'Calibri'

# ══════════════════════════════════════════
# EDUCATION + CERTIFICATIONS + LANGUAGES (combined row)
# ══════════════════════════════════════════
section_head("Education")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = Pt(11)
r = p.add_run("Bachelor of Business & Information Technology (BBIT)")
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'
r2 = p.add_run(" - Virtual University of Pakistan | CGPA: 3.85/4.00 | Expected June 2029")
r2.font.size = Pt(9)
r2.font.color.rgb = GREY
r2.font.name = 'Calibri'

section_head("Certifications & Awards")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = Pt(11)
r = p.add_run("Harvard CS50x & Aspire Leader Program | Data Science Professional (KNIME) | Deep Learning Essentials (UOP) | Business & Data Analysis (Microsoft/LinkedIn) | 2026 Young Leaders Fellowship Nominee (Rotaract Bahrain)")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
r.font.name = 'Calibri'

section_head("Languages")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("English (Fluent) | Urdu (Native)")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
r.font.name = 'Calibri'

# Save
output_dir = r"f:\Uzmentum"
os.makedirs(output_dir, exist_ok=True)
word_path = os.path.join(output_dir, "Muhammad_Usman_Resume_2026.docx")
doc.save(word_path)
print(f"Word resume saved to: {word_path}")